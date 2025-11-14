"""Document processing utilities."""
import os
from typing import Optional
from pathlib import Path


class DocumentProcessor:
    """Process various document formats."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document

            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from any supported document format."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif ext == ".txt":
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def save_as_docx(content: str, output_path: str, title: Optional[str] = None):
        """Save content as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()

            # Add title if provided
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = 1  # Center alignment

            # Add content (split by paragraphs)
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())

            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Failed to save DOCX: {str(e)}")

    @staticmethod
    def save_as_txt(content: str, output_path: str):
        """Save content as text file."""
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            return output_path
        except Exception as e:
            raise Exception(f"Failed to save TXT: {str(e)}")
